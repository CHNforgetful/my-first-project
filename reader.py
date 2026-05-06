import requests
import json
import os
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

# 填你的API Key
API_KEY = "sk-c8cffb5896eb4332b92b055d1f0e00fe"  

# DeepSeek最新模型配置
MODEL_NAME = "deepseek-v4-flash"  # 可选: deepseek-v4-flash 或 deepseek-v4-pro
BASE_URL = "https://api.deepseek.com"

def get_new_urls():
    """读取urls.txt，返回未处理的URL列表"""
    
    # 读取所有URL
    with open("urls.txt", "r", encoding="utf-8") as f:
        all_urls = [line.strip() for line in f if line.strip()]
    
    # 读取已处理的URL
    processed = set()
    if os.path.exists("processed.txt"):
        with open("processed.txt", "r", encoding="utf-8") as f:
            processed = set(line.strip() for line in f)
    
    # 返回新URL
    new_urls = [url for url in all_urls if url not in processed]
    return new_urls, processed

def save_processed_url(url):
    """保存已处理的URL"""
    with open("processed.txt", "a", encoding="utf-8") as f:
        f.write(url + "\n")

def fetch_article(url):
    """抓取文章内容"""
    # 没搞明白

    from bs4 import BeautifulSoup
    
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
    }
    
    try:
        response = requests.get(url, headers=headers, timeout=15)
        response.raise_for_status()
        
        soup = BeautifulSoup(response.text, 'html.parser')
        
        # 获取标题
        title_tag = soup.find('title')
        title = title_tag.text.strip() if title_tag else "无标题"
        
        # 获取正文
        paragraphs = soup.find_all('p')
        content = '\n\n'.join([p.text for p in paragraphs])
        
        # 清理多余空格
        content = ' '.join(content.split())
        
        return title, content
        
    except Exception as e:
        print(f"   抓取错误: {e}")
        return None, None

def analyze_article(article):
    """分析英文文章（带错误处理）"""

    # 固定部分（会被缓存）
    system_prompt = """你是英语阅读学习助手。
分析英文文章，用中文回复，格式如下：

【摘要】50字左右的中文摘要

【词组】5-7个，如动词短语、固定搭配或习惯用语等，不选简单词组
词组：中文释义
原文例句

【好句】2-4句，要句式优美或包含重要语法
原文句子
中文翻译

规则：
1.翻译要准确通顺
2.词组例句和好句不能重复。如果某个句子同时适合两者，优先放在【好句】中
3.不要输出任何额外说明或评价
"""
    
    # 变化部分
    user_prompt = f"文章：\n{article[:3000]}"

    try:
        response = requests.post(
            f"{BASE_URL}/v1/chat/completions", 
            headers={
                "Authorization": f"Bearer {API_KEY}",
                "Content-Type": "application/json"
            },
        json={
            "model": MODEL_NAME,
            "messages": [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            "temperature": 0.3,
            "thinking": {"type": "disabled"}
        },
            timeout=30  # 设置超时
        )
        
        # 检查HTTP状态码
        if response.status_code != 200:
            return f"❌ API错误：{response.status_code}\n{response.text}"
        
        # 解析JSON
        result = response.json()
        
        # 提取内容
        if "choices" in result and len(result["choices"]) > 0:
            return result["choices"][0]["message"]["content"]
        else:
            return f"❌ 返回格式异常：{result}"
            
    except requests.exceptions.Timeout:
        return "❌ 请求超时，请重试"
    except requests.exceptions.ConnectionError:
        return "❌ 网络连接失败，请检查网络"
    except json.JSONDecodeError as e:
        return f"❌ JSON解析失败：{e}\n原始返回：{response.text[:200]}"
    except Exception as e:
        return f"❌ 未知错误：{e}"

def save_to_word(title, analysis, url):
    """保存为Word"""
    # 没搞懂
    Path("阅读报告").mkdir(exist_ok=True)
    
    doc = Document()
    
#     # 直接写入内容，不做任何格式处理
#     content = f"""标题：{title}
# 原文链接：{url}
# 分析时间：{datetime.now().strftime("%Y-%m-%d %H:%M:%S")}
# {'=' * 50}

# {analysis}
# """
    
#     # 按行添加，保持原有格式
#     for line in content.split('\n'):
#         doc.add_paragraph(line)
    
        # 标题
    p = doc.add_paragraph(f'标题：{title}')
    if p.runs:
        set_font(p.runs[0], 16)
    
    # 链接
    p = doc.add_paragraph(f'原文链接：{url}')
    if p.runs:
        set_font(p.runs[0], 10)
    
    # 时间
    p = doc.add_paragraph(f'分析时间：{datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    if p.runs:
        set_font(p.runs[0], 10)
    
    # 分隔线
    p = doc.add_paragraph('=' * 50)
    if p.runs:
        set_font(p.runs[0], 10)
    
    # 正文（处理空行）
    for line in analysis.split('\n'):
        p = doc.add_paragraph(line)
        if p.runs and line.strip():  # 只有非空行才设置字体
            set_font(p.runs[0], 12)

    # 保存
    safe_title = "".join(c for c in title[:30] if c.isalnum() or c in ' _')
    filename = f"{datetime.now().strftime('%Y%m%d_%H%M%S')}_{safe_title}.docx"
    filepath = Path("阅读报告") / filename
    doc.save(filepath)
    
    return filepath

def set_font(run, size, font_name='微软雅黑'):
    """设置字体大小和类型"""
    run.font.size = Pt(size)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def main():
    print("=" * 60)
    print("英语阅读批量处理助手")
    print("=" * 60)
    
    # 1. 获取新URL
    new_urls, processed = get_new_urls()
    
    if not new_urls:
        print("\n没有新文章需要处理！")
        print("\n使用方法：")
        print("   1. 在 urls.txt 中添加文章URL（每行一个）")
        print("   2. 重新运行本程序")
        return
    
    print(f"\n📌 发现 {len(new_urls)} 篇新文章\n")
    
    # 2. 逐个处理
    for i, url in enumerate(new_urls, 1):
        print(f"[{i}/{len(new_urls)}] 处理中...")
        print(f"   URL: {url[:60]}...")
        
        # 抓取文章
        print("   抓取文章...")
        title, content = fetch_article(url)
        
        if not title or not content:
            print("  抓取失败，跳过\n")
            save_processed_url(url)
            continue
        
        print(f"   标题: {title[:40]}")
        print(f"   字数: {len(content)}")
        
        # AI分析
        print("   AI分析中...")
        try:
            analysis = analyze_article(content)
        except Exception as e:
            print(f"   分析失败: {e}\n")
            save_processed_url(url)
            continue
        
        # 保存结果
        print("   保存结果...")
        result_file = save_to_word(title, analysis, url)
        
        # 标记已处理
        save_processed_url(url)
        
        print(f"   完成！{result_file}\n")
    
    print("=" * 60)
    print(f"全部完成！共处理 {len(new_urls)} 篇文章")
    print(f"结果保存在 reports/ 文件夹")
    print("=" * 60)

if __name__ == "__main__":
    main()
